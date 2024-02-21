import docx
import pandas as pd
import re

import warnings
warnings.filterwarnings('ignore')

Input_excel_file = 'Input.xlsx'
ExcelData_df = pd.read_excel(Input_excel_file)
def detect_hyperlink_columns(df):

    hyperlink_columns = []
    for col in df.columns:
        for value in df[col]:
            if isinstance(value, str) and re.match(r'https?://\S+', value):
                hyperlink_columns.append(col)
                break  
    return hyperlink_columns

def add_hyperlink(cell, url, color, underline):
    paragraph = cell.paragraphs[0]
    paragraph.clear() 
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)
    new_run.append(rPr)
    new_run.text = url 
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink



hyperlink_columns = detect_hyperlink_columns(ExcelData_df)

if not hyperlink_columns:
    print("No URL-like columns detected in the DataFrame.")
else:
    Word = docx.Document()
    table = Word.add_table(rows=1, cols=len(ExcelData_df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(ExcelData_df.columns):
        hdr_cells[i].text = col_name
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True


    for _, row in ExcelData_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

            for cell in row_cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

        # Add hyperlinks if the current column matches any of the detected hyperlink columns
        for col in hyperlink_columns:
            cell = row_cells[ExcelData_df.columns.get_loc(col)]
            hyperlink = add_hyperlink(cell, row[col], '0000FF', False)

    # Save the Word document
    Word.save('Output.docx')
    print("Done , file saved as Output.docx at cwd ")
